# app.py   ——  支持 Word/Excel 双版本请款单
import os, re, zipfile, io, tempfile
from pathlib import Path
from datetime import date
import pandas as pd
from docx import Document
from openpyxl import load_workbook
import streamlit as st

# ---------- 小工具 ----------
def number_to_upper(amount: int) -> str:
    CN_NUM = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    CN_UNIT = ['', '拾', '佰', '仟', '万', '拾万', '佰万', '仟万', '亿']
    if amount == 0:
        return "零元整"
    s = str(amount)
    result = []
    for i, ch in enumerate(s[::-1]):
        digit = int(ch)
        if digit != 0:
            result.append(f"{CN_NUM[digit]}{CN_UNIT[i]}")
    upper_str = ''.join(reversed(result))
    upper_str = re.sub(r'零{2,}', '零', upper_str)
    upper_str = re.sub(r'零元', '元', upper_str)
    upper_str = re.sub(r'亿万', '亿', upper_str)
    return upper_str + "元整"

def sanitize_filename(name: str) -> str:
    return re.sub(r'[<>:\"/\\|?*\x00-\x1F]', '_', name).rstrip('. ')

# ---------- Excel 请款单生成 ----------
def build_excel_invoice(split_no, sub_df, template_path, output_dir, company_name):
    applicant = str(sub_df["申请人"].iloc[0]) if "申请人" in sub_df.columns else ""
    official_total = pd.to_numeric(sub_df["官费"], errors="coerce").fillna(0).sum()
    agent_total    = pd.to_numeric(sub_df["代理费"], errors="coerce").fillna(0).sum()
    grand_total    = official_total + agent_total

    # 案号顿号拼接
    case_no_list = []
    for _, row in sub_df.iterrows():
        if "集佳案号" in sub_df.columns and pd.notna(row.get("集佳案号")):
            case_no_list.append(str(row["集佳案号"]))
        elif "我方案号" in sub_df.columns and pd.notna(row.get("我方案号")):
            case_no_list.append(str(row["我方案号"]))
    case_no_str = "、".join(case_no_list)

    wb = load_workbook(template_path)
    ws = wb.active

    # 占位符替换
    for row in ws.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            v = str(cell.value)
            v = v.replace("{{合计}}", str(int(grand_total))) \
                 .replace("{{日期}}", date.today().strftime("%Y年%m月%d日")) \
                 .replace("{{申请人}}", applicant)
            cell.value = v

    # 数据写入：从第3行开始覆盖
    start_row = 3
    for _, r in sub_df.iterrows():
        for col_idx, col_name in enumerate(sub_df.columns, 1):
            ws.cell(row=start_row, column=col_idx, value=r[col_name])
        start_row += 1

    # 文件名
    filename = sanitize_filename(f"请款单（{applicant}-{int(grand_total)}元-{company_name}-{date.today().strftime('%Y-%m-%d')}).xlsx")
    wb.save(output_dir / filename)
    print(f"✅ 已生成Excel请款单：{filename}")
    return {
        "分割号": split_no,
        "申请人": applicant,
        "总官费": int(official_total),
        "总代理费": int(agent_total),
        "总计": int(grand_total),
        "文件名": filename
    }

# ---------- Word 请款单生成 ----------
def build_word_invoice(split_no, sub_df, template_path, output_dir, company_name):
    applicant = str(sub_df["申请人"].iloc[0]) if "申请人" in sub_df.columns else ""
    official_total = pd.to_numeric(sub_df["官费"], errors="coerce").fillna(0).sum()
    agent_total    = pd.to_numeric(sub_df["代理费"], errors="coerce").fillna(0).sum()
    grand_total    = official_total + agent_total

    case_no_list = []
    for _, row in sub_df.iterrows():
        if "集佳案号" in sub_df.columns and pd.notna(row.get("集佳案号")):
            case_no_list.append(str(row["集佳案号"]))
        elif "我方案号" in sub_df.columns and pd.notna(row.get("我方案号")):
            case_no_list.append(str(row["我方案号"]))
    case_no_str = "、".join(case_no_list)

    doc = Document(template_path)

    # 占位符段落替换（保留格式）
    placeholders = {
        "{{申请人}}": applicant,
        "{{合计}}": str(int(grand_total)),
        "{{大写}}": number_to_upper(int(grand_total)),
        "{{日期}}": date.today().strftime("%Y年%m月%d日")
    }
    for p in doc.paragraphs:
        for key, val in placeholders.items():
            for run in p.runs:
                if key in run.text:
                    run.text = run.text.replace(key, val)

    # 表格填充
    tbl = doc.tables[0]
    hdr_cells = tbl.rows[0].cells
    for idx, col_name in enumerate(sub_df.columns):
        if idx < len(hdr_cells):
            hdr_cells[idx].text = str(col_name)

    for _, r in sub_df.iterrows():
        new_cells = tbl.add_row().cells
        for idx, col_name in enumerate(sub_df.columns):
            if idx < len(new_cells):
                new_cells[idx].text = str(r[col_name] or "")

    # 合计行
    try:
        off_idx = sub_df.columns.get_loc("官费")
        agt_idx = sub_df.columns.get_loc("代理费")
    except KeyError:
        off_idx, agt_idx = 0, 1
    row = tbl.add_row()
    cells = row.cells
    merge_start = cells[0]
    merge_end = cells[off_idx - 1] if off_idx > 0 else cells[0]
    if merge_start != merge_end:
        merge_start.merge(merge_end)
    merge_start.text = "合计"
    cells[off_idx].text = str(int(official_total))
    cells[agt_idx].text = str(int(agent_total))
    cells[agt_idx + 1].text = str(int(grand_total))

    filename = sanitize_filename(f"请款单（{applicant}-{int(grand_total)}元-{company_name}-{date.today().strftime('%Y-%m-%d')}).docx")
    doc.save(output_dir / filename)
    print(f"✅ 已生成Word请款单：{filename}")
    return {
        "分割号": split_no,
        "申请人": applicant,
        "总官费": int(official_total),
        "总代理费": int(agent_total),
        "总计": int(grand_total),
        "文件名": filename
    }

# ---------- Streamlit UI ----------
def main():
    st.set_page_config(page_title="专利请款单生成器", page_icon="📄", layout="centered")
    st.markdown("""
    <style>
    .main-header{font-size:2.2rem;color:#0066cc;text-align:center;margin-bottom:1rem;}
    .upload-card{background:#fff;border:1px solid #cce0ff;border-radius:12px;padding:1.5rem 2rem;margin-bottom:1.5rem;box-shadow:0 2px 6px rgba(0,102,204,.08);}
    .stButton>button{background:#4da6ff;color:#fff;border:none;padding:.45rem 1.4rem;border-radius:8px}
    .stButton>button:hover{background:#0077e6}
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<h1 class="main-header">📄 专利请款单生成器</h1>', unsafe_allow_html=True)

    # 1. 选择版本
    st.markdown('<div class="upload-card">', unsafe_allow_html=True)
    version = st.radio("请选择生成版本", ("Word 请款单", "Excel 请款单"), horizontal=True)
    st.markdown('</div>', unsafe_allow_html=True)

    # 2. 上传模板
    st.markdown('<div class="upload-card">', unsafe_allow_html=True)
    st.subheader("📤 上传模板 & 数据")
    col1, col2 = st.columns(2)
    with col1:
        tpl = st.file_uploader(
            "上传请款单模板",
            type=["docx"] if version == "Word 请款单" else ["xlsx"],
            help="请确保模板已包含占位符：{{合计}}、{{日期}}、{{申请人}}"
        )
    with col2:
        data = st.file_uploader(
            "上传需请款专利清单",
            type=["xlsx"],
            help="清单必须包含“分割号、官费、代理费、申请人”列"
        )
    st.markdown('</div>', unsafe_allow_html=True)

    # 3. 公司选择
    st.markdown('<div class="upload-card" style="text-align:center;">', unsafe_allow_html=True)
    company_name = st.radio("命名格式", ["深佳", "集佳"], horizontal=True, label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)

    # 4. 生成按钮
    col_btn = st.columns([1, 2, 1])
    with col_btn[1]:
        if st.button("🚀 生成文件", use_container_width=True, type="primary"):
            if not tpl or not data:
                st.error("请上传模板和数据！")
                st.stop()

            with tempfile.TemporaryDirectory() as tmp:
                tmp_path = Path(tmp)
                out_path = tmp_path / "output"
                out_path.mkdir()

                tpl_path = tmp_path / ("tpl.docx" if version.startswith("Word") else "tpl.xlsx")
                data_path = tmp_path / "data.xlsx"
                with open(tpl_path, "wb") as f:
                    f.write(tpl.getbuffer())
                with open(data_path, "wb") as f:
                    f.write(data.getbuffer())

                df = pd.read_excel(data_path, dtype=str).fillna("")
                if any(c not in df.columns for c in ["分割号", "官费", "代理费"]):
                    st.error("数据清单缺少必要列！")
                    st.stop()

                rows, succ, err = [], 0, 0
                bar = st.progress(0)
                total = len(df.groupby("分割号"))

                for idx, (split_no, sub) in enumerate(df.groupby("分割号")):
                    try:
                        if version.startswith("Word"):
                            res = build_word_invoice(split_no, sub, tpl_path, out_path, company_name)
                        else:
                            res = build_excel_invoice(split_no, sub, tpl_path, out_path, company_name)
                        rows.append(res)
                        succ += 1
                    except Exception as e:
                        err += 1
                        st.warning(f"⚠ 分割号 {split_no} 失败：{e}")
                    bar.progress((idx + 1) / total)

                if rows:
                    try:
                        generate_invoice_excel(rows, out_path, Path(tmp_path) / "发票申请表.xlsx", company_name)
                    except Exception:
                        pass

                # 收集下载
                if 'files' not in st.session_state:
                    st.session_state['files'] = {}
                st.session_state['files'] = {
                    f.name: f.read_bytes() for f in out_path.iterdir()
                }
                st.success(f"🎉 完成！成功 {succ}，失败 {err}")

    # 5. 下载区
    if 'files' in st.session_state and st.session_state['files']:
        st.markdown("---")
        st.subheader("📥 下载生成的文件")
        col_zip = st.columns([1, 2, 1])
        with col_zip[1]:
            if st.button("📦 一键打包", use_container_width=True, type="secondary"):
                buf = io.BytesIO()
                with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
                    for fn, data_bytes in st.session_state['files'].items():
                        zf.writestr(fn, data_bytes)
                buf.seek(0)
                st.download_button("⬇️ 下载 ZIP", buf,
                                   f"请款单_{company_name}_{date.today():%Y%m%d}.zip",
                                   "application/zip", use_container_width=True)

        col1, col2 = st.columns(2)
        for col, ext, mime, label in [(col1, "docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "Word"),
                                      (col2, "xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "Excel")]:
            with col:
                st.write(f"**{label}文件:**")
                files = {k: v for k, v in st.session_state['files'].items() if k.endswith(ext)}
                if files:
                    for fn, data_bytes in files.items():
                        st.download_button(f"下载 {fn}", data_bytes, fn, mime, use_container_width=True)
                else:
                    st.info(f"暂无{label}文件")

if __name__ == "__main__":
    main()
